#!/usr/bin/env python3
"""Debug script to analyze footnote extraction and placement"""

import sys
sys.path.append('/home/shaul/Scripts')
from docx_to_org_complete import DocxToOrgCompleteConverter
from docx import Document

def debug_footnotes(docx_file):
    print(f"=== Debugging footnotes in {docx_file} ===")
    
    converter = DocxToOrgCompleteConverter()
    doc = Document(docx_file)
    
    # Extract footnotes
    converter.extract_footnotes_from_xml(doc)
    print(f"Found {len(converter.footnotes)} footnotes:")
    for i, fn in enumerate(converter.footnotes, 1):
        print(f"  [fn:{i}] ID={fn['id']}, Text: {fn['text'][:50]}...")
    
    # Analyze each paragraph for footnote references
    print("\n=== Analyzing paragraphs for footnote references ===")
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            refs = converter.find_footnote_references_in_paragraph(paragraph)
            if refs:
                print(f"Paragraph {para_idx+1}: Found {len(refs)} references")
                for run_idx, ref in refs:
                    print(f"  Run {run_idx}: {ref}")
                print(f"  Text: {paragraph.text[:100]}...")
    
    # Check for potential footnote references in text
    print("\n=== Looking for potential footnote patterns in text ===")
    import re
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        # Look for superscript numbers or symbols
        if text.strip():
            # Look for patterns that might be footnote references
            potential_refs = re.findall(r'[\u00b9\u00b2\u00b3]+|[\u2070-\u2079]+|†+|‡+|§+', text)
            if potential_refs:
                print(f"Paragraph {para_idx+1}: Potential refs {potential_refs}")
                print(f"  Text: {text}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        debug_footnotes(sys.argv[1])
    else:
        debug_footnotes("/home/shaul/Documents/siach_shai/b_typed/a_siach_shai/a_siach_shai/a_siach_shai_shas/a_siach_shai/f_taharos/nidda/nidda_beginning.docx")