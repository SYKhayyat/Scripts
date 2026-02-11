#!/usr/bin/env python3
"""Debug script to check exactly what happens during conversion of specific paragraphs"""

import sys
sys.path.append('/home/shaul/Scripts')
from docx_to_org_complete import DocxToOrgCompleteConverter
from docx import Document

def debug_specific_paragraphs(docx_file):
    print(f"=== Debugging specific paragraphs in {docx_file} ===")
    
    converter = DocxToOrgCompleteConverter()
    doc = Document(docx_file)
    
    # Extract footnotes
    converter.extract_footnotes_from_xml(doc)
    print(f"Found {len(converter.footnotes)} footnotes")
    
    # Check specific paragraphs that should have footnotes 1, 2, 5
    target_paragraphs = [8, 14, 16]  # 0-indexed for paragraphs 9, 15, 17
    
    for para_idx in target_paragraphs:
        if para_idx < len(doc.paragraphs):
            paragraph = doc.paragraphs[para_idx]
            if paragraph.text.strip():
                print(f"\n--- Paragraph {para_idx + 1} ---")
                print(f"Original text (first 200 chars): {paragraph.text[:200]}...")
                
                # Find footnote references
                refs = converter.find_footnote_references_in_paragraph(paragraph)
                print(f"Found footnote references: {refs}")
                
                # Check if has bold
                has_bold = converter.has_bold_text(paragraph)
                print(f"Has bold text: {has_bold}")
                
                # Process with footnotes using the current method
                processed_text = converter._process_text_with_footnotes(paragraph)
                print(f"Processed text (first 200 chars): {processed_text[:200]}...")
                
                # Check if footnotes are in the processed text
                for _, ref in refs:
                    if ref in processed_text:
                        print(f"✅ {ref} found in processed text")
                    else:
                        print(f"❌ {ref} MISSING from processed text!")
                
                # If it has bold, also check the bold extraction method
                if has_bold:
                    bold, regular = converter.extract_bold_portions_with_footnotes(paragraph)
                    print(f"Bold portions: {len(bold)}")
                    print(f"Regular portions: {len(regular)}")
                    for i, b in enumerate(bold):
                        if any(ref in b for _, ref in refs):
                            print(f"  Bold {i}: contains footnote")
                    for i, r in enumerate(regular):
                        if any(ref in r for _, ref in refs):
                            print(f"  Regular {i}: contains footnote")

if __name__ == "__main__":
    debug_specific_paragraphs("/home/shaul/Documents/siach_shai/b_typed/a_siach_shai/a_siach_shai/a_siach_shai_shas/a_siach_shai/f_taharos/nidda/nidda_beginning.docx")