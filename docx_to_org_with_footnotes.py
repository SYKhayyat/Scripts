#!/usr/bin/env python3
"""
Enhanced Docx to Org-mode Converter with Footnote Support
Converts Hebrew docx files to org-mode format with proper header and footnote handling.
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Tuple, Optional, Dict
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx library not installed. Please run:")
    print("pip install python-docx")
    sys.exit(1)


class DocxToOrgWithFootnotesConverter:
    def __init__(self):
        self.footnotes = []
        self.footnote_counter = 1
        self.footnote_references = {}  # Maps original reference to footnote number
    
    def get_user_files(self) -> List[Tuple[str, str]]:
        """Interactive file selection for input and output paths"""
        file_pairs = []
        
        print("=== Enhanced Docx to Org-mode Converter with Footnotes ===")
        print("Enter file paths (or 'done' to finish, 'quit' to exit)")
        
        while True:
            # Get input file
            input_file = input("\nEnter docx file path: ").strip()
            
            if input_file.lower() == 'quit':
                return []
            elif input_file.lower() == 'done':
                break
            elif not input_file:
                print("Please enter a valid file path")
                continue
            
            # Check if file exists
            if not os.path.exists(input_file):
                print(f"File not found: {input_file}")
                continue
            
            # Check if it's a docx file
            if not input_file.lower().endswith('.docx'):
                print("File must be a .docx file")
                continue
            
            # Suggest output file
            suggested_output = input_file.replace('.docx', '.org')
            output_file = input(f"Enter output file path [{suggested_output}]: ").strip()
            
            if not output_file:
                output_file = suggested_output
            
            file_pairs.append((input_file, output_file))
            print(f"Added: {input_file} -> {output_file}")
        
        return file_pairs
    
    def is_centered(self, paragraph) -> bool:
        """Check if paragraph is centered"""
        try:
            return paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        except:
            return False
    
    def has_bold_text(self, paragraph) -> bool:
        """Check if paragraph contains bold text"""
        for run in paragraph.runs:
            if run.bold:
                return True
        return False
    
    def extract_bold_portions(self, paragraph) -> Tuple[List[str], List[str]]:
        """Extract bold and regular text portions from a paragraph"""
        bold_text = []
        regular_text = []
        current_bold = []
        current_regular = []
        
        for run in paragraph.runs:
            if run.bold:
                if current_regular:
                    regular_text.append(''.join(current_regular))
                    current_regular = []
                current_bold.append(run.text)
            else:
                if current_bold:
                    bold_text.append(''.join(current_bold))
                    current_bold = []
                current_regular.append(run.text)
        
        # Add any remaining text
        if current_bold:
            bold_text.append(''.join(current_bold))
        if current_regular:
            regular_text.append(''.join(current_regular))
        
        return bold_text, regular_text
    
    def extract_footnotes_from_xml(self, doc) -> None:
        """Extract footnotes from document XML"""
        try:
            # Method 1: Try to access footnotes through the document part
            if hasattr(doc, '_part'):
                document_part = doc._part
                
                # Try to get footnotes part
                try:
                    footnotes_part = document_part.package.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
                    )
                    
                    if footnotes_part:
                        # Parse the footnotes XML
                        footnotes_xml = footnotes_part.blob
                        if footnotes_xml:
                            import xml.etree.ElementTree as ET
                            root = ET.fromstring(footnotes_xml)
                            
                            # Define namespace
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            
                            # Find all footnote elements
                            footnote_elements = root.findall('.//w:footnote', ns)
                            
                            for footnote in footnote_elements:
                                # Get footnote ID
                                footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                if footnote_id and footnote_id != '0':  # Skip separator/continuation notes
                                    # Extract footnote text
                                    footnote_text = self._extract_text_from_footnote_element(footnote)
                                    if footnote_text.strip():
                                        self.footnotes.append({
                                            'id': footnote_id,
                                            'text': footnote_text.strip()
                                        })
                except Exception as e:
                    print(f"Warning: Could not access footnotes part: {e}")
                
                # Try to get endnotes part
                try:
                    endnotes_part = document_part.package.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
                    )
                    
                    if endnotes_part:
                        # Parse the endnotes XML
                        endnotes_xml = endnotes_part.blob
                        if endnotes_xml:
                            import xml.etree.ElementTree as ET
                            root = ET.fromstring(endnotes_xml)
                            
                            # Define namespace
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            
                            # Find all endnote elements
                            endnote_elements = root.findall('.//w:endnote', ns)
                            
                            for endnote in endnote_elements:
                                # Get endnote ID
                                endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                if endnote_id and endnote_id != '0':  # Skip separator/continuation notes
                                    # Extract endnote text
                                    endnote_text = self._extract_text_from_footnote_element(endnote)
                                    if endnote_text.strip():
                                        self.footnotes.append({
                                            'id': endnote_id,
                                            'text': endnote_text.strip()
                                        })
                except Exception as e:
                    print(f"Warning: Could not access endnotes part: {e}")
            
            # Method 2: Try alternative approach using document relationships
            try:
                if hasattr(doc, '_part',) and hasattr(doc._part, 'rels'):
                    # Look for footnotes relationship
                    for rel_id, rel in doc._part.rels.items():
                        if hasattr(rel, 'reltype') and rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes':
                            footnotes_part = rel.target_part
                            if footnotes_part:
                                import xml.etree.ElementTree as ET
                                root = ET.fromstring(footnotes_part.blob)
                                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                footnote_elements = root.findall('.//w:footnote', ns)
                                
                                for footnote in footnote_elements:
                                    footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                    if footnote_id and footnote_id != '0':
                                        footnote_text = self._extract_text_from_footnote_element(footnote)
                                        if footnote_text.strip():
                                            self.footnotes.append({
                                                'id': footnote_id,
                                                'text': footnote_text.strip()
                                            })
                        
                        # Check for endnotes
                        if hasattr(rel, 'reltype') and rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes':
                            endnotes_part = rel.target_part
                            if endnotes_part:
                                import xml.etree.ElementTree as ET
                                root = ET.fromstring(endnotes_part.blob)
                                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                endnote_elements = root.findall('.//w:endnote', ns)
                                
                                for endnote in endnote_elements:
                                    endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                    if endnote_id and endnote_id != '0':
                                        endnote_text = self._extract_text_from_footnote_element(endnote)
                                        if endnote_text.strip():
                                            self.footnotes.append({
                                                'id': endnote_id,
                                                'text': endnote_text.strip()
                                            })
            except Exception as e:
                print(f"Warning: Alternative footnote extraction failed: {e}")
                    
        except Exception as e:
            print(f"Warning: Could not extract footnotes from XML: {e}")
    
    def _extract_text_from_footnote_element(self, footnote_element) -> str:
        """Extract text content from a footnote XML element"""
        try:
            text_parts = []
            # Get all text runs within the footnote
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elements = footnote_element.findall('.//w:t', ns)
            for text_elem in text_elements:
                if text_elem.text:
                    text_parts.append(text_elem.text)
            return ''.join(text_parts)
        except:
            return ""
    
    def detect_footnote_references_in_text(self, text: str) -> List[int]:
        """Detect potential footnote references in text (numbers, symbols, etc.)"""
        import re
        # Look for superscript-like patterns: numbers, symbols
        patterns = [
            r'[\d]+',  # numbers
            r'[†‡§¶]',  # common footnote symbols
        ]
        
        references = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                try:
                    references.append(int(match))
                except ValueError:
                    # For symbols, just add them as is
                    references.append(match)
        
        return references
    
    def find_footnote_references_in_paragraph(self, paragraph) -> List[Tuple[str, str]]:
        """Find footnote references within a paragraph and return (original_text, replacement) pairs"""
        references = []
        
        try:
            # Check each run for footnote references
            for run_idx, run in enumerate(paragraph.runs):
                # Look for footnote reference elements in the run's XML
                if hasattr(run, 'element'):
                    run_xml = run.element
                    # Check for footnote reference elements
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    footnote_refs = run_xml.findall('.//w:footnoteReference', ns)
                    
                    for ref in footnote_refs:
                        ref_id = ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        if ref_id:
                            # Find the corresponding footnote
                            footnote_num = None
                            for i, fn in enumerate(self.footnotes):
                                if fn['id'] == ref_id:
                                    footnote_num = i + 1
                                    break
                            
                            if footnote_num:
                                # Create replacement for the footnote reference
                                replacement = f"[fn:{footnote_num}]"
                                # We'll use the run index to identify where to place the reference
                                references.append((run_idx, replacement))
        except Exception as e:
            print(f"Warning: Error finding footnote references: {e}")
        
        return references
    
    def _process_text_with_footnotes(self, paragraph) -> str:
        """Process paragraph text and insert footnote references at appropriate positions"""
        footnote_refs = self.find_footnote_references_in_paragraph(paragraph)
        
        if not footnote_refs:
            return paragraph.text
        
        # Build text with footnote references
        result_parts = []
        processed_runs = set()
        
        for run_idx, run in enumerate(paragraph.runs):
            # Check if this run has a footnote reference
            run_footnotes = [ref for ref in footnote_refs if ref[0] == run_idx]
            
            if run_footnotes:
                # Add the run text followed by footnote references
                if run.text:
                    result_parts.append(run.text)
                for _, footnote_ref in run_footnotes:
                    result_parts.append(footnote_ref)
                processed_runs.add(run_idx)
            else:
                # Just add the run text
                if run.text:
                    result_parts.append(run.text)
        
        return ''.join(result_parts)
    
    def _apply_footnote_references_to_text(self, text: str, footnote_refs, paragraph) -> str:
        """Apply footnote references to extracted text portions"""
        if not footnote_refs:
            return text
        
        # For now, append footnote references at the end of text portions
        # This is a simplified approach - in practice, we'd need more sophisticated positioning
        result = text
        for run_idx, footnote_ref in footnote_refs:
            # Try to find if this text corresponds to a run with footnotes
            if run_idx < len(paragraph.runs) and paragraph.runs[run_idx].text in text:
                # If the text matches the run content, append the footnote reference
                if not result.endswith(footnote_ref):
                    result += footnote_ref
        
        return result
    
    def process_paragraph(self, paragraph) -> List[str]:
        """Process a single paragraph and return org-mode lines with footnote references"""
        lines = []
        
        # Get the processed text with footnotes
        processed_text = self._process_text_with_footnotes(paragraph)
        
        if not processed_text.strip():
            return lines
        
        # Handle centered text as level 1 header
        if self.is_centered(paragraph):
            lines.append(f"* {processed_text.strip()}")
            return lines
        
        # Handle bold text in non-centered paragraphs
        if self.has_bold_text(paragraph):
            bold_portions, regular_portions = self.extract_bold_portions(paragraph)
            
            # Add bold portions as level 2 headers
            for bold_text in bold_portions:
                if bold_text.strip():
                    lines.append(f"** {bold_text.strip()}")
            
            # Add regular portions as normal paragraphs
            for regular_text in regular_portions:
                if regular_text.strip():
                    lines.append(regular_text.strip())
        else:
            # Regular paragraph with footnote processing
            lines.append(processed_text.strip())
        
        return lines
    
    def process_paragraph_with_footnotes(self, paragraph) -> str:
        """Simplified method to process paragraph with footnotes"""
        # Get the full text with footnote references
        processed_text = self._process_text_with_footnotes(paragraph)
        return processed_text.strip()
    
    def convert_docx_to_org(self, input_file: str, output_file: str) -> bool:
        """Convert a single docx file to org-mode with footnotes"""
        try:
            print(f"Converting {input_file}...")
            
            # Reset footnote data
            self.footnotes = []
            self.footnote_counter = 1
            self.footnote_references = {}
            
            # Load document
            doc = Document(input_file)
            
            # Extract footnotes first
            self.extract_footnotes_from_xml(doc)
            
            if self.footnotes:
                print(f"Found {len(self.footnotes)} footnotes")
            else:
                print("No footnotes found in document")
                # Check for potential footnote references in text
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        refs = self.detect_footnote_references_in_text(para.text)
                        if refs:
                            print(f"Paragraph {i+1} contains potential footnote references: {refs}")
            
            org_content = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                lines = self.process_paragraph(paragraph)
                org_content.extend(lines)
            
            # Add footnotes at the end if any
            if self.footnotes:
                org_content.append("")
                org_content.append("* Footnotes")
                for i, footnote in enumerate(self.footnotes, 1):
                    org_content.append(f"[fn:{i}] {footnote['text']}")
            
            # Write to org file with UTF-8 encoding
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(org_content))
            
            print(f"✓ Converted to {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ Error converting {input_file}: {str(e)}")
            return False
    
    def run(self):
        """Main conversion loop"""
        file_pairs = self.get_user_files()
        
        if not file_pairs:
            print("No files selected. Exiting.")
            return
        
        print(f"\nProcessing {len(file_pairs)} file(s)...")
        
        success_count = 0
        for input_file, output_file in file_pairs:
            if self.convert_docx_to_org(input_file, output_file):
                success_count += 1
        
        print(f"\nConversion complete: {success_count}/{len(file_pairs)} files processed successfully.")


if __name__ == "__main__":
    converter = DocxToOrgWithFootnotesConverter()
    
    # Check if command line arguments are provided
    if len(sys.argv) >= 2:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '.org')
        
        if os.path.exists(input_file) and input_file.lower().endswith('.docx'):
            if converter.convert_docx_to_org(input_file, output_file):
                print(f"Successfully converted {input_file} to {output_file}")
            else:
                print(f"Failed to convert {input_file}")
        else:
            print(f"Invalid input file: {input_file}")
    else:
        converter.run()