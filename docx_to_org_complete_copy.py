#!/usr/bin/env python3
"""
Complete Docx to Org-mode Converter with Footnotes and List Support
Converts Hebrew docx files to org-mode format with proper header, footnote, and list handling.
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Tuple, Optional, Dict
import glob
from concurrent.futures import ThreadPoolExecutor, as_completed
from multiprocessing import cpu_count
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx library not installed. Please run:")
    print("pip install python-docx")
    sys.exit(1)


class DocxToOrgCompleteConverter:
    def __init__(self):
        self.footnotes = []
        self.footnote_counter = 1
        self.footnote_references = {}  # Maps original reference to footnote number
        self.list_depth_stack = []  # Track nested list levels
        self.list_item_counters = []  # Track numbering for ordered lists
        self.global_action = None  # Store global conflict resolution action
    
    def find_docx_files_recursive(self, directory: str) -> List[str]:
        """Find all .docx files in directory recursively"""
        docx_files = []
        try:
            for root, dirs, files in os.walk(directory):
                for file in files:
                    if file.lower().endswith('.docx'):
                        docx_files.append(os.path.join(root, file))
        except Exception as e:
            print(f"Error scanning directory {directory}: {e}")
        return docx_files
    
    def ask_user_for_conflict_resolution(self, output_file: str):
        """Ask user how to resolve file conflict, returns (should_process, action, global_action)"""
        choice_map = {
            '1': (True, 'overwrite', None),
            '2': (False, 'skip', None),
            '3': (True, 'append_single', None),
            '4': (True, 'overwrite', 'overwrite_all'),
            '5': (False, 'skip', 'skip_all'),
            '6': (True, 'append', 'append_all')
        }
        
        while True:
            print(f"File '{output_file}' already exists. Choose an option:")
            print("  1) Overwrite this file")
            print("  2) Skip this file") 
            print("  3) Add (1) to this filename")
            print("  4) Overwrite all files")
            print("  5) Skip all files")
            print("  6) Add (1) to all conflicted files")
            
            choice = input("Enter choice (1-6): ").strip()
            
            if choice in choice_map:
                return choice_map[choice]
            else:
                print("Please enter a number between 1 and 6")
    
    def resolve_conflict(self, output_file: str, conflict_mode: str = "ask", global_action: Optional[str] = None):
        """Resolve file conflicts based on mode, returns (should_process, final_filename, global_action)"""
        if conflict_mode == "overwrite":
            return True, output_file, global_action
        elif conflict_mode == "suffix":
            return True, self.add_numeric_suffix(output_file), global_action
        elif conflict_mode == "append":
            return True, output_file.replace('.org', '(1).org') if output_file.endswith('.org') else output_file + '(1)', global_action
        elif conflict_mode == "ask":
            if global_action:
                # Use the global action if already determined - no more asking
                if global_action == 'overwrite_all':
                    return True, output_file, global_action
                elif global_action == 'skip_all':
                    return False, output_file, global_action
                elif global_action == 'append_all':
                    final_output = output_file.replace('.org', '(1).org') if output_file.endswith('.org') else output_file + '(1)'
                    return True, final_output, global_action
            else:
                # First conflict - ask user for resolution
                should_process, action, new_global = self.ask_user_for_conflict_resolution(output_file)
                
                if action == 'overwrite':
                    return True, output_file, new_global  # Individual action, no global change
                elif action == 'skip':
                    return False, output_file, new_global   # Individual action, no global change
                elif action == 'append_single':
                    final_output = output_file.replace('.org', '(1).org') if output_file.endswith('.org') else output_file + '(1)'
                    return True, final_output, new_global   # Individual action, no global change
                elif action == 'overwrite':
                    return True, output_file, new_global  # Should not happen, but handle
                elif action == 'append':
                    final_output = output_file.replace('.org', '(1).org') if output_file.endswith('.org') else output_file + '(1)'
                    return True, final_output, new_global  # Global action set
                else:
                    return should_process, output_file, new_global
        else:
            print(f"Unknown conflict mode: {conflict_mode}")
            should_process, action, new_global = self.ask_user_for_conflict_resolution(output_file)
            if action == 'append_single' or action == 'append':
                final_output = output_file.replace('.org', '(1).org') if output_file.endswith('.org') else output_file + '(1)'
                return should_process, final_output, new_global
            return should_process, output_file, new_global
    
    def add_numeric_suffix(self, output_file: str) -> str:
        """Add numeric suffix to filename if file exists"""
        base_name = output_file
        suffix = 1
        
        while os.path.exists(base_name):
            # Remove existing suffix if any
            if suffix > 1:
                base_name = base_name.rstrip(f".{suffix-1}")
            
            new_name = f"{base_name}.{suffix}"
            base_name = new_name
            suffix += 1
        
        return base_name
    
    def get_user_input(self) -> List[Tuple[str, str]]:
        """Get input from user - can be file or folder"""
        inputs = []
        
        print("=== Complete Docx to Org-mode Converter (Footnotes + Lists + Batch) ===")
        print("Enter file or folder paths (or 'done' to finish, 'quit' to exit)")
        print("Folders will be scanned recursively for .docx files")
        
        while True:
            input_path = input("\nEnter docx file or folder path: ").strip()
            
            if input_path.lower() == 'quit':
                return []
            elif input_path.lower() == 'done':
                break
            elif not input_path:
                print("Please enter a valid path")
                continue
            
            # Check if path exists
            if not os.path.exists(input_path):
                print(f"Path not found: {input_path}")
                continue
            
            # Handle folder
            if os.path.isdir(input_path):
                docx_files = self.find_docx_files_recursive(input_path)
                if not docx_files:
                    print(f"No .docx files found in {input_path}")
                    continue
                
                print(f"Found {len(docx_files)} .docx files in {input_path}")
                confirm = input(f"Convert all {len(docx_files)} files? [y/n] ").strip().lower()
                if confirm in ['y', 'yes']:
                    for docx_file in docx_files:
                        output_file = docx_file.replace('.docx', '.org')
                        inputs.append((docx_file, output_file))
                else:
                    continue
            
            # Handle single file
            elif input_path.lower().endswith('.docx'):
                suggested_output = input_path.replace('.docx', '.org')
                output_file = input(f"Enter output file path [{suggested_output}]: ").strip()
                if not output_file:
                    output_file = suggested_output
                inputs.append((input_path, output_file))
            
            else:
                print("Path must be a .docx file or folder")
                continue
            
            print(f"Added: {len(inputs)} file(s) queued for conversion")
        
        return inputs
    
    def get_user_files(self) -> List[Tuple[str, str]]:
        """Interactive file selection for input and output paths"""
        file_pairs = []
        
        print("=== Complete Docx to Org-mode Converter (Footnotes + Lists) ===")
        print("Enter file paths (or 'done' to finish, 'quit' to exit)")
        
        while True:
            # Get input file
            input_file = input("\nEnter docx file path: ").strip()
            
            if input_file.lower() == 'quit':
                return []
            elif input_file.lower() == 'done':
                break
            elif not input_file:
                print("Please enter a valid file path")
                continue
            
            # Check if file exists
            if not os.path.exists(input_file):
                print(f"File not found: {input_file}")
                continue
            
            # Check if it's a docx file
            if not input_file.lower().endswith('.docx'):
                print("File must be a .docx file")
                continue
            
            # Suggest output file
            suggested_output = input_file.replace('.docx', '.org')
            output_file = input(f"Enter output file path [{suggested_output}]: ").strip()
            
            if not output_file:
                output_file = suggested_output
            
            file_pairs.append((input_file, output_file))
            print(f"Added: {input_file} -> {output_file}")
        
        return file_pairs
    
    def is_centered(self, paragraph) -> bool:
        """Check if paragraph is centered"""
        try:
            return paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        except:
            return False
    
    def has_bold_text(self, paragraph) -> bool:
        """Check if paragraph contains bold text"""
        for run in paragraph.runs:
            if run.bold:
                return True
        return False
    
    def extract_bold_portions(self, paragraph) -> Tuple[List[str], List[str]]:
        """Extract bold and regular text portions from a paragraph"""
        bold_text = []
        regular_text = []
        current_bold = []
        current_regular = []
        
        for run in paragraph.runs:
            if run.bold:
                if current_regular:
                    regular_text.append(''.join(current_regular))
                    current_regular = []
                current_bold.append(run.text)
            else:
                if current_bold:
                    bold_text.append(''.join(current_bold))
                    current_bold = []
                current_regular.append(run.text)
        
        # Add any remaining text
        if current_bold:
            bold_text.append(''.join(current_bold))
        if current_regular:
            regular_text.append(''.join(current_regular))
        
        return bold_text, regular_text
    
    def extract_bold_portions_with_footnotes(self, paragraph) -> Tuple[List[str], List[str]]:
        """Extract bold and regular text portions while preserving footnote references"""
        bold_text = []
        regular_text = []
        current_bold = []
        current_regular = []
        
        # Find footnote references first
        footnote_refs = self.find_footnote_references_in_paragraph(paragraph)
        footnote_ref_map = {run_idx: ref for run_idx, ref in footnote_refs}
        
        for run_idx, run in enumerate(paragraph.runs):
            # Get the text for this run
            run_text = run.text
            
            # Add footnote reference if this run has one
            if run_idx in footnote_ref_map:
                run_text += footnote_ref_map[run_idx]
            
            if run.bold:
                if current_regular:
                    regular_text.append(''.join(current_regular))
                    current_regular = []
                current_bold.append(run_text)
            else:
                if current_bold:
                    bold_text.append(''.join(current_bold))
                    current_bold = []
                current_regular.append(run_text)
        
        # Add any remaining text
        if current_bold:
            bold_text.append(''.join(current_bold))
        if current_regular:
            regular_text.append(''.join(current_regular))
        
        return bold_text, regular_text
    
    def detect_list_type_and_level(self, paragraph) -> Tuple[Optional[str], int]:
        """Detect if paragraph is a list item and determine list type and level"""
        try:
            # Method 1: Check paragraph style
            if hasattr(paragraph, 'style') and paragraph.style:
                style_name = paragraph.style.name.lower()
                if 'list' in style_name:
                    # Determine list type from style name
                    if 'bullet' in style_name or 'unordered' in style_name:
                        # Try to get level from style XML
                        ilvl_elements = paragraph.style.element.xpath('.//w:ilvl')
                        level_val = int(ilvl_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if ilvl_elements else 0
                        
                        # If XML level is 0, try fallback from style name
                        if level_val == 0:
                            # Fallback: try to extract from style name like "list bullet 2"
                            if style_name.endswith(' 2'):
                                level_val = 1
                            elif style_name.endswith(' 3'):
                                level_val = 2
                        return 'unordered', level_val
                    elif 'number' in style_name or 'ordered' in style_name:
                        # Try to get level from style XML
                        ilvl_elements = paragraph.style.element.xpath('.//w:ilvl')
                        level_val = int(ilvl_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if ilvl_elements else 0
                        
                        # If XML level is 0, try fallback from style name
                        if level_val == 0:
                            # Fallback: try to extract from style name like "list number 2"
                            if style_name.endswith(' 2'):
                                level_val = 1
                            elif style_name.endswith(' 3'):
                                level_val = 2
                        return 'ordered', level_val
            
            # Method 2: Check paragraph numbering properties
            if hasattr(paragraph, '_element'):
                element = paragraph._element
                # Look for numbering elements
                num_elements = element.xpath('.//w:numPr')
                if num_elements:
                    num_pr = num_elements[0]
                    # Get indentation level
                    ilvl_elements = num_pr.xpath('.//w:ilvl')
                    level = ilvl_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0') if ilvl_elements else '0'
                    
                    # Get numbering ID to determine list type
                    num_id_elements = num_pr.xpath('.//w:numId')
                    if num_id_elements:
                        return 'ordered', int(level)
                    else:
                        return 'unordered', int(level)
            
            # Method 3: Check text patterns (fallback)
            text = paragraph.text.strip()
            if text:
                # Check for bullet patterns
                bullet_patterns = [
                    r'^[•·▪▫‣⁃]\s+',
                    r'^[o*+−-]\s+',
                    r'^[◦◉○●]\s+'
                ]
                
                for pattern in bullet_patterns:
                    if re.match(pattern, text):
                        return 'unordered', 0
                
                # Check for numbered patterns
                numbered_patterns = [
                    r'^\d+\.\s+',
                    r'^\d+\)\s+',
                    r'^[a-zA-Z]\.\s+',
                    r'^[a-zA-Z]\)\s+',
                    r'^[ivxlIVXL]+\.\s+',
                    r'^[IVXL]+\)\s+'
                ]
                
                for pattern in numbered_patterns:
                    if re.match(pattern, text):
                        return 'ordered', 0
                
                # Check for indented text (might be nested list)
                if text.startswith('    ') or text.startswith('\t'):
                    return 'unordered', 1
            
        except Exception as e:
            print(f"Warning: Error detecting list type: {e}")
        
        return None, 0
    
    def clean_list_text(self, text: str, list_type: str) -> str:
        """Remove list markers from text and clean it up"""
        text = text.strip()
        
        if list_type == 'unordered':
            # Remove common bullet markers
            bullet_patterns = [
                r'^[•·▪▫‣⁃]\s+',
                r'^[o*+−-]\s+',
                r'^[◦◉○●]\s+',
                r'^[•]\s+'
            ]
            for pattern in bullet_patterns:
                text = re.sub(pattern, '', text)
        elif list_type == 'ordered':
            # Remove numbered markers
            numbered_patterns = [
                r'^\d+\.\s+',
                r'^\d+\)\s+',
                r'^[a-zA-Z]\.\s+',
                r'^[a-zA-Z]\)\s+',
                r'^[ivxlIVXL]+\.\s+',
                r'^[IVXL]+\)\s+'
            ]
            for pattern in numbered_patterns:
                text = re.sub(pattern, '', text)
        
        return text.strip()
    
    def get_org_list_prefix(self, list_type: str, level: int, item_number: Optional[int] = None) -> str:
        """Get appropriate org-mode list prefix based on type and level"""
        indent = '  ' * level
        
        if list_type == 'unordered':
            # Alternate between -, +, and * at different levels
            bullets = ['-', '+', '*']
            bullet = bullets[level % len(bullets)]
            return f"{indent}{bullet} "
        elif list_type == 'ordered':
            # Org-mode uses 1., 2., 3. for ordered lists at all levels
            if item_number is None:
                item_number = 1
            return f"{indent}{item_number}. "
        
        return ''
    
    def extract_footnotes_from_xml(self, doc) -> None:
        """Extract footnotes from document XML"""
        try:
            # Method 1: Try to access footnotes through the document part
            if hasattr(doc, '_part'):
                document_part = doc._part
                
                # Try to get footnotes part
                try:
                    footnotes_part = document_part.package.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
                    )
                    
                    if footnotes_part:
                        # Parse the footnotes XML
                        footnotes_xml = footnotes_part.blob
                        if footnotes_xml:
                            import xml.etree.ElementTree as ET
                            root = ET.fromstring(footnotes_xml)
                            
                            # Define namespace
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            
                            # Find all footnote elements
                            footnote_elements = root.findall('.//w:footnote', ns)
                            
                            for footnote in footnote_elements:
                                # Get footnote ID
                                footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                if footnote_id and footnote_id != '0':  # Skip separator/continuation notes
                                    # Extract footnote text
                                    footnote_text = self._extract_text_from_footnote_element(footnote)
                                    if footnote_text.strip():
                                        self.footnotes.append({
                                            'id': footnote_id,
                                            'text': footnote_text.strip()
                                        })
                except Exception as e:
                    print(f"Warning: Could not access footnotes part: {e}")
                
                # Try to get endnotes part
                try:
                    endnotes_part = document_part.package.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
                    )
                    
                    if endnotes_part:
                        # Parse the endnotes XML
                        endnotes_xml = endnotes_part.blob
                        if endnotes_xml:
                            import xml.etree.ElementTree as ET
                            root = ET.fromstring(endnotes_xml)
                            
                            # Define namespace
                            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            
                            # Find all endnote elements
                            endnote_elements = root.findall('.//w:endnote', ns)
                            
                            for endnote in endnote_elements:
                                # Get endnote ID
                                endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                if endnote_id and endnote_id != '0':  # Skip separator/continuation notes
                                    # Extract endnote text
                                    endnote_text = self._extract_text_from_footnote_element(endnote)
                                    if endnote_text.strip():
                                        self.footnotes.append({
                                            'id': endnote_id,
                                            'text': endnote_text.strip()
                                        })
                except Exception as e:
                    print(f"Warning: Could not access endnotes part: {e}")
            
            # Method 2: Try alternative approach using document relationships
            try:
                if hasattr(doc, '_part',) and hasattr(doc._part, 'rels'):
                    # Look for footnotes relationship
                    for rel_id, rel in doc._part.rels.items():
                        if hasattr(rel, 'reltype') and rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes':
                            footnotes_part = rel.target_part
                            if footnotes_part:
                                import xml.etree.ElementTree as ET
                                root = ET.fromstring(footnotes_part.blob)
                                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                footnote_elements = root.findall('.//w:footnote', ns)
                                
                                for footnote in footnote_elements:
                                    footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                    if footnote_id and footnote_id != '0':
                                        footnote_text = self._extract_text_from_footnote_element(footnote)
                                        if footnote_text.strip():
                                            self.footnotes.append({
                                                'id': footnote_id,
                                                'text': footnote_text.strip()
                                            })
                        
                        # Check for endnotes
                        if hasattr(rel, 'reltype') and rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes':
                            endnotes_part = rel.target_part
                            if endnotes_part:
                                import xml.etree.ElementTree as ET
                                root = ET.fromstring(endnotes_part.blob)
                                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                endnote_elements = root.findall('.//w:endnote', ns)
                                
                                for endnote in endnote_elements:
                                    endnote_id = endnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                                    if endnote_id and endnote_id != '0':
                                        endnote_text = self._extract_text_from_footnote_element(endnote)
                                        if endnote_text.strip():
                                            self.footnotes.append({
                                                'id': endnote_id,
                                                'text': endnote_text.strip()
                                            })
            except Exception as e:
                print(f"Warning: Alternative footnote extraction failed: {e}")
                    
        except Exception as e:
            print(f"Warning: Could not extract footnotes from XML: {e}")
    
    def _extract_text_from_footnote_element(self, footnote_element) -> str:
        """Extract text content from a footnote XML element"""
        try:
            text_parts = []
            # Get all text runs within the footnote
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elements = footnote_element.findall('.//w:t', ns)
            for text_elem in text_elements:
                if text_elem.text:
                    text_parts.append(text_elem.text)
            return ''.join(text_parts)
        except:
            return ""
    
    def detect_footnote_references_in_text(self, text: str) -> List[int]:
        """Detect potential footnote references in text (numbers, symbols, etc.)"""
        import re
        # Look for superscript-like patterns: numbers, symbols
        patterns = [
            r'\d+',  # numbers
            r'[†‡§¶]',  # common footnote symbols
        ]
        
        references = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                try:
                    references.append(int(match))
                except ValueError:
                    # For symbols, just add them as is
                    references.append(match)
        
        return references
    
    def find_footnote_references_in_paragraph(self, paragraph) -> List[Tuple[str, str]]:
        """Find footnote references within a paragraph and return (original_text, replacement) pairs"""
        references = []
        
        try:
            # Check each run for footnote references
            for run_idx, run in enumerate(paragraph.runs):
                # Look for footnote reference elements in the run's XML
                if hasattr(run, 'element'):
                    run_xml = run.element
                    # Check for footnote reference elements
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    footnote_refs = run_xml.findall('.//w:footnoteReference', ns)
                    
                    for ref in footnote_refs:
                        ref_id = ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        if ref_id:
                            # Find the corresponding footnote
                            footnote_num = None
                            for i, fn in enumerate(self.footnotes):
                                if fn['id'] == ref_id:
                                    footnote_num = i + 1
                                    break
                            
                            if footnote_num:
                                # Create replacement for the footnote reference
                                replacement = f"[fn:{footnote_num}]"
                                # We'll use the run index to identify where to place the reference
                                references.append((run_idx, replacement))
        except Exception as e:
            print(f"Warning: Error finding footnote references: {e}")
        
        return references
    
    def _process_text_with_footnotes(self, paragraph) -> str:
        """Process paragraph text and insert footnote references at appropriate positions"""
        footnote_refs = self.find_footnote_references_in_paragraph(paragraph)
        
        if not footnote_refs:
            return paragraph.text
        
        # Build text with footnote references
        result_parts = []
        processed_runs = set()
        
        for run_idx, run in enumerate(paragraph.runs):
            # Check if this run has a footnote reference
            run_footnotes = [ref for ref in footnote_refs if ref[0] == run_idx]
            
            if run_footnotes:
                # Add the run text followed by footnote references
                if run.text:
                    result_parts.append(run.text)
                for _, footnote_ref in run_footnotes:
                    result_parts.append(footnote_ref)
                processed_runs.add(run_idx)
            else:
                # Just add the run text
                if run.text:
                    result_parts.append(run.text)
        
        return ''.join(result_parts)
    
    def _apply_footnote_references_to_text(self, text: str, footnote_refs, paragraph) -> str:
        """Apply footnote references to extracted text portions"""
        if not footnote_refs:
            return text
        
        # For now, append footnote references at the end of text portions
        # This is a simplified approach - in practice, we'd need more sophisticated positioning
        result = text
        for run_idx, footnote_ref in footnote_refs:
            # Try to find if this text corresponds to a run with footnotes
            if run_idx < len(paragraph.runs) and paragraph.runs[run_idx].text in text:
                # If the text matches the run content, append the footnote reference
                if not result.endswith(footnote_ref):
                    result += footnote_ref
        
        return result
    
    def process_paragraph(self, paragraph) -> List[str]:
        """Process a single paragraph and return org-mode lines with footnote references and list support"""
        lines = []
        
        # Get the processed text with footnotes
        processed_text = self._process_text_with_footnotes(paragraph)
        
        if not processed_text.strip():
            return lines
        
        # Check if this is a list item
        list_type, level = self.detect_list_type_and_level(paragraph)
        
        if list_type:
            # Handle list item
            cleaned_text = self.clean_list_text(processed_text, list_type)
            
            # Update list depth stack
            while len(self.list_depth_stack) > level + 1:
                self.list_depth_stack.pop()
                self.list_item_counters.pop()
            
            while len(self.list_depth_stack) < level + 1:
                self.list_depth_stack.append(list_type)
                self.list_item_counters.append(1)
            
            # Get item number for ordered lists
            item_number = None
            if list_type == 'ordered':
                item_number = self.list_item_counters[level]
                self.list_item_counters[level] += 1
            else:
                # Reset ordered list counter when we encounter unordered list at same level
                if level < len(self.list_item_counters):
                    self.list_item_counters[level] = 1
            
            # Create org-mode list item
            prefix = self.get_org_list_prefix(list_type, level, item_number if list_type == 'ordered' else 1)
            lines.append(f"{prefix}{cleaned_text}")
            
        else:
            # Reset list stack when we encounter non-list content
            self.list_depth_stack = []
            self.list_item_counters = []
            
            # Handle centered text as level 1 header
            if self.is_centered(paragraph):
                lines.append(f"* {processed_text.strip()}")
                return lines
            
            # Handle bold text in non-centered paragraphs
            if self.has_bold_text(paragraph):
                # Use the enhanced method that preserves footnotes
                bold_portions, regular_portions = self.extract_bold_portions_with_footnotes(paragraph)
                
                # Add bold portions as level 2 headers
                for bold_text in bold_portions:
                    if bold_text.strip():
                        lines.append(f"** {bold_text.strip()}")
                
                # Add regular portions as normal paragraphs
                for regular_text in regular_portions:
                    if regular_text.strip():
                        lines.append(regular_text.strip())
            else:
                # Regular paragraph with footnote processing
                lines.append(processed_text.strip())
        
        return lines
    
    def process_paragraph_with_footnotes(self, paragraph) -> str:
        """Simplified method to process paragraph with footnotes"""
        # Get the full text with footnote references
        processed_text = self._process_text_with_footnotes(paragraph)
        return processed_text.strip()
    
    def convert_docx_to_org(self, input_file: str, output_file: str, check_conflicts: bool = True, conflict_mode: str = "ask") -> bool:
        """Convert a single docx file to org-mode with footnotes and lists"""
        try:
            # Handle conflicts based on mode
            if check_conflicts and os.path.exists(output_file):
                result = self.resolve_conflict(output_file, conflict_mode, self.global_action)
                if result is None:
                    print(f"⏭ Skipped {input_file} (conflict resolution failed)")
                    return False
                should_process, final_output, new_global_action = result
                
                # Store the global action for future conflicts
                if new_global_action:
                    self.global_action = new_global_action
                
                if not should_process:
                    print(f"⏭ Skipped {input_file} (conflict resolved by user)")
                    return False
                output_file = final_output
            
            print(f"Converting {input_file}...")
            
            # Reset all data
            self.footnotes = []
            self.footnote_counter = 1
            self.footnote_references = {}
            self.list_depth_stack = []
            self.list_item_counters = []
            
            # Load document
            doc = Document(input_file)
            
            # Extract footnotes first
            self.extract_footnotes_from_xml(doc)
            
            if self.footnotes:
                print(f"Found {len(self.footnotes)} footnotes")
            else:
                print("No footnotes found in document")
                # Check for potential footnote references in text
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        refs = self.detect_footnote_references_in_text(para.text)
                        if refs:
                            print(f"Paragraph {i+1} contains potential footnote references: {refs}")
            
            org_content = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                lines = self.process_paragraph(paragraph)
                org_content.extend(lines)
            
            # Add footnotes at the end if any
            if self.footnotes:
                org_content.append("")
                org_content.append("* Footnotes")
                for i, footnote in enumerate(self.footnotes, 1):
                    org_content.append(f"[fn:{i}] {footnote['text']}")
            
            # Create output directory if it doesn't exist
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Write to org file with UTF-8 encoding
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(org_content))
            
            print(f"✓ Converted to {output_file}")
            return True
            
        except Exception as e:
            print(f"✗ Error converting {input_file}: {str(e)}")
            return False
    
    @staticmethod
    def convert_single_file_static(input_file: str, output_file: str, check_conflicts: bool = True, conflict_mode: str = "ask") -> Tuple[str, bool]:
        """Convert a single file for parallel processing, returns (filename, success)"""
        # Create a new converter instance for thread safety
        converter = DocxToOrgCompleteConverter()
        success = converter.convert_docx_to_org(input_file, output_file, check_conflicts, conflict_mode)
        return (output_file, success)
    
    def convert_single_file(self, input_file: str, output_file: str, check_conflicts: bool = True, conflict_mode: str = "ask", global_action: Optional[str] = None) -> Tuple[str, bool]:
        """Convert a single file for parallel processing, returns (filename, success)"""
        # Set the global action for this conversion
        if global_action:
            self.global_action = global_action
        success = self.convert_docx_to_org(input_file, output_file, check_conflicts, conflict_mode)
        return (output_file, success)
    
    def run(self, parallel: bool = True, max_workers: Optional[int] = None):
        """Main conversion loop with optional parallel processing"""
        self.global_action = None  # Reset global action for this run
        file_pairs = self.get_user_input()
        
        if not file_pairs:
            print("No files selected. Exiting.")
            return
        
        print(f"\nProcessing {len(file_pairs)} file(s)...")
        if parallel and len(file_pairs) > 1:
            max_workers = max_workers or min(cpu_count(), len(file_pairs))
            print(f"Using parallel processing with {max_workers} workers")
            
            success_count = 0
            skip_count = 0
            
            # For parallel processing, pre-resolve conflicts to avoid asking during conversion
            resolved_file_pairs = []
            for input_file, output_file in file_pairs:
                if os.path.exists(output_file):
                    # Resolve conflict now before parallel processing
                    result = self.resolve_conflict(output_file, "ask", self.global_action)
                    if result:
                        should_process, final_output, new_global_action = result
                        
                        # Store the global action for future conflicts
                        if new_global_action:
                            self.global_action = new_global_action
                        
                        if should_process:
                            resolved_file_pairs.append((input_file, final_output))
                        else:
                            print(f"⏭ Skipped {input_file} (conflict resolved)")
                            skip_count += 1
                    else:
                        print(f"⏭ Skipped {input_file} (conflict resolution failed)")
                        skip_count += 1
                else:
                    resolved_file_pairs.append((input_file, output_file))
            
            success_count = 0
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all conversion tasks with pre-resolved conflicts
                future_to_file = {
                    executor.submit(DocxToOrgCompleteConverter.convert_single_file_static, input_file, output_file, False, "overwrite"): (input_file, output_file)
                    for input_file, output_file in resolved_file_pairs
                }
                
                # Process completed tasks
                for future in as_completed(future_to_file):
                    input_file, output_file = future_to_file[future]
                    try:
                        result_file, success = future.result()
                        if success:
                            success_count += 1
                        else:
                            skip_count += 1
                    except Exception as e:
                        print(f"✗ Error processing {input_file}: {str(e)}")
                        skip_count += 1
        else:
            # Sequential processing
            success_count = 0
            skip_count = 0
            for input_file, output_file in file_pairs:
                if self.convert_docx_to_org(input_file, output_file):
                    success_count += 1
                else:
                    skip_count += 1
        
        print(f"\nConversion complete: {success_count} succeeded, {skip_count} skipped out of {len(file_pairs)} files.")


if __name__ == "__main__":
    converter = DocxToOrgCompleteConverter()
    
    # Check if command line arguments are provided
    if len(sys.argv) >= 2:
        input_path = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        conflict_mode = "ask"  # default mode
        overwrite_all = False
        
        # Parse additional arguments
        parallel = True
        max_workers = None
        for arg in sys.argv[3:]:
            if arg.lower() in ['--overwrite-all', '-o']:
                overwrite_all = True
                conflict_mode = "overwrite"
            elif arg.lower() in ['--add-suffix', '-a']:
                overwrite_all = False
                conflict_mode = "suffix"
            elif arg.lower() in ['--append', '-p']:
                overwrite_all = False
                conflict_mode = "append"
            elif arg.lower() == '--no-parallel':
                parallel = False
            elif arg.lower().startswith('--workers='):
                try:
                    max_workers = int(arg.split('=', 1)[1])
                except ValueError:
                    print(f"Invalid worker count: {arg}")
                    sys.exit(1)
            else:
                print(f"Unknown argument: {arg}")
                print("Usage: python3 docx_to_org_complete_copy.py <input> [output] [--overwrite-all|-o] [--add-suffix|-a] [--append|-p] [--no-parallel] [--workers=N]")
                print("  --overwrite-all, -o: Overwrite all existing files")
                print("  --add-suffix, -a: Add numeric suffix to conflicting files")
                print("  --append, -p: Append (1) to conflicting files")
                print("  --no-parallel: Disable parallel processing")
                print("  --workers=N: Set number of parallel workers")
                sys.exit(1)
        
        if not os.path.exists(input_path):
            print(f"Path not found: {input_path}")
            sys.exit(1)
        
        # Handle single file
        if os.path.isfile(input_path) and input_path.lower().endswith('.docx'):
            if not output_file:
                output_file = input_path.replace('.docx', '.org')
            if converter.convert_docx_to_org(input_path, output_file, check_conflicts=False, conflict_mode=conflict_mode):
                print(f"Successfully converted {input_path} to {output_file}")
            else:
                print(f"Failed to convert {input_path}")
        
        # Handle folder
        elif os.path.isdir(input_path):
            docx_files = converter.find_docx_files_recursive(input_path)
            if not docx_files:
                print(f"No .docx files found in {input_path}")
                sys.exit(1)
            
            print(f"Found {len(docx_files)} .docx files in {input_path}")
            
            if parallel and len(docx_files) > 1:
                max_workers = max_workers or min(cpu_count(), len(docx_files))
                print(f"Using parallel processing with {max_workers} workers")
                
                success_count = 0
                file_pairs = [(docx_file, docx_file.replace('.docx', '.org')) for docx_file in docx_files]
                
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    future_to_file = {
                        executor.submit(converter.convert_single_file, input_file, output_file, False, conflict_mode): (input_file, output_file)
                        for input_file, output_file in file_pairs
                    }
                    
                    for future in as_completed(future_to_file):
                        input_file, output_file = future_to_file[future]
                        try:
                            result_file, success = future.result()
                            if success:
                                success_count += 1
                        except Exception as e:
                            print(f"✗ Error processing {input_file}: {str(e)}")
            else:
                success_count = 0
                for docx_file in docx_files:
                    org_file = docx_file.replace('.docx', '.org')
                    if converter.convert_docx_to_org(docx_file, org_file, check_conflicts=False, conflict_mode=conflict_mode):
                        success_count += 1
            
            print(f"Batch conversion complete: {success_count}/{len(docx_files)} files converted successfully.")
        
        else:
            print(f"Invalid input: {input_path} (must be .docx file or folder)")
            sys.exit(1)
    else:
        converter.run()